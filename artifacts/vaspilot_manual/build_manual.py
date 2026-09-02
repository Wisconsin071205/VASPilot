from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "VASPilot_详细使用手册.docx"
UI = ROOT / "artifacts" / "vaspilot_manual" / "ui"
VLAB = ROOT / "artifacts" / "vaspilot_manual" / "vlab"
CODEX_DASHBOARD = Path(
    r"C:\Users\weikx\AppData\Local\Temp\codex-clipboard-45f76d47-f9bd-49c5-8fb2-976324d9dc25.png"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "237A57"
AMBER = "8A6200"
RED = "A1352B"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei",
                 size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(paragraph, color: str, size: int = 18,
                          space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = paragraph.add_run("第 ")
    set_run_font(r1, size=9, color=MUTED)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=9, color=MUTED)
    r2 = paragraph.add_run(" 页")
    set_run_font(r2, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(32, 40, 48)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.paragraph_format.space_after = Pt(16)

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.10)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05


def patch_numbering(doc: Document) -> None:
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        for lvl in abstract.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) not in {
                "bullet", "decimal"
            }:
                continue
            p_pr = lvl.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                lvl.append(p_pr)
            tabs = p_pr.find(qn("w:tabs"))
            if tabs is None:
                tabs = OxmlElement("w:tabs")
                p_pr.append(tabs)
            for child in list(tabs):
                tabs.remove(child)
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), "269")
            tabs.append(tab)
            ind = p_pr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                p_pr.append(ind)
            ind.set(qn("w:left"), "540")
            ind.set(qn("w:hanging"), "271")


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def new_decimal_num_id(doc: Document) -> int:
    """Create a fresh decimal numbering instance that starts at 1."""
    numbering = doc.part.numbering_part.element
    style_num_pr = doc.styles["List Number"]._element.pPr.numPr
    base_num_id = str(style_num_pr.numId.val)
    base_num = next(
        (node for node in numbering.findall(qn("w:num"))
         if node.get(qn("w:numId")) == base_num_id),
        None,
    )
    if base_num is None:
        raise RuntimeError("List Number style has no numbering definition")
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return new_id


def add_numbers(doc: Document, items: list[str]) -> None:
    num_id = new_decimal_num_id(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
        r = p.add_run(item)
        set_run_font(r)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    shade_paragraph(p, "EEF2F6")
    paragraph_left_border(p, BLUE, size=12, space=6)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", east_asia="Microsoft YaHei",
                     size=9, color="243142")


def add_note(doc: Document, label: str, text: str,
             kind: str = "info") -> None:
    color = {"info": BLUE, "warn": AMBER, "danger": RED, "ok": GREEN}[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    shade_paragraph(p, CALLOUT)
    paragraph_left_border(p, color, size=22, space=8)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="29333D")


def add_h1(doc: Document, text: str, page_break=True) -> None:
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.page_break_before = page_break
    set_keep_with_next(p)


def add_h2(doc: Document, text: str, page_break=False) -> None:
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.page_break_before = page_break
    set_keep_with_next(p)


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 3")
    set_keep_with_next(p)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[int], compact: bool = False) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if idx == 0
                       else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=(8.8 if compact else 9.5), color=NAVY, bold=True)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if idx == 0
                           else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = (1.0 if compact else 1.08)
            r = p.add_run(str(value))
            set_run_font(r, size=(8.5 if compact else 9.2))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, path: Path, caption: str, alt: str,
               width=6.35) -> None:
    if not path.is_file():
        add_note(doc, "界面图缺失", f"未找到截图：{path}", "warn")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = True


def add_kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(value)
    set_run_font(r, size=10.5)


def add_section_overview(doc: Document, items: list[tuple[str, str]]) -> None:
    for label, value in items:
        add_kv(doc, label, value)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("VASPilot 使用手册")
    set_run_font(hr, size=9, color=MUTED, bold=True)
    hr = hp.add_run("    |    Codex · Web UI · CLI")
    set_run_font(hr, size=9, color=MUTED)
    footer = section.footer
    add_page_field(footer.paragraphs[0])

    # Cover — editorial_cover pattern, compact_reference_guide body.
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("HPC / VASP 智能计算平台")
    set_run_font(r, size=11, color=BLUE, bold=True)
    title = doc.add_paragraph("VASPilot 详细使用手册", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Codex 插件、本地 Web 控制台与命令行操作指南", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向胡伟团队研究人员与平台维护人员")
    set_run_font(r, size=11, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(32)
    add_figure(
        doc, CODEX_DASHBOARD,
        "封面图：Codex 中的远程集群作业面板（状态数据会随时间变化）",
        "VASPilot Codex fleet dashboard showing multiple HPC server cards",
        width=6.15,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(18)
    r = meta.add_run("适用源码界面：VASPilot 1.3.5\n")
    set_run_font(r, size=10, color=NAVY, bold=True)
    r = meta.add_run("Codex 已安装 Client 后端：1.3.0\n")
    set_run_font(r, size=10, color=MUTED)
    r = meta.add_run(f"编制日期：{date.today().isoformat()}")
    set_run_font(r, size=10, color=MUTED)

    add_h1(doc, "阅读说明", page_break=True)
    add_section_overview(doc, [
        ("适用对象", "第一次接触 VASPilot 的研究人员、日常使用人员，以及负责服务器与模型配置的维护人员"),
        ("覆盖入口", "Codex 内嵌集群面板、本地 Web 控制台、Windows 命令行"),
        ("使用目标", "完成服务器连接、文件查看、作业监测、VASP 科学进度判断、工作流计划与模型对话"),
        ("界面版本", "截图取自本地源码 UI 1.3.5；Codex 插件当前优先加载的 Client 后端为 1.3.0，个别按钮或修复可能略有差异"),
    ])
    add_note(doc, "重要概念",
             "“作业完成”只代表调度器生命周期结束；是否达到 VASP 电子或离子收敛，必须另查“VASP 科学进度”。",
             "warn")
    add_h2(doc, "内容导航")
    add_bullets(doc, [
        "系统全貌：理解本机、Vlab、HPC、调度器和 VASP 解析之间的关系",
        "快速开始：首次启动、连接服务器、打开 Codex 面板",
        "Web UI：对话、文件、作业、工作流、监测、配置六大页面",
        "CLI：服务器、远程文件、作业、工作流、监测和模型命令",
        "端到端范例：日常巡检、已有计算诊断、完整提交和多服务器选择",
        "故障排查：登录、主机指纹、网关、调度器、VASP、模型与 UI",
        "附录：状态词典、常用提示词、交接检查表",
    ])

    add_h1(doc, "1. 系统全貌")
    add_body(doc, "VASPilot 的核心不是让模型直接登录服务器，而是把本地界面、模型调用、Vlab 跳板机和多台 HPC 服务器串成一条可审计的数据通路。日常用户主要面对 Codex、Web UI 或 CLI；SSH 复用、调度器差异和 VASP 输出解析由后端处理。")
    add_code(doc, "Codex / Web UI / CLI\n"
                  "        ↓ 共享工具注册表\n"
                  "本机 GatewayClient 与审计日志\n"
                  "        ↓ SSH\n"
                  "USTC Vlab 网关\n"
                  "        ↓ 每服务器 SSH 复用会话\n"
                  "Slurm / PBS 集群 → VASP 输入、作业、输出与科学进度")
    add_h2(doc, "1.1 三类状态必须分开")
    add_table(doc, ["状态维度", "回答的问题", "典型字段", "结论边界"], [
        ["连接状态", "现在能否访问服务器？", "connected、auth_mode、reconnect_state", "只说明 SSH 通路"],
        ["调度器状态", "作业在排队、运行还是结束？", "PENDING、RUNNING、COMPLETED、FAILED", "不说明科学收敛"],
        ["科学状态", "VASP 是否电子/离子收敛？", "scientific_converged、ionic_steps、NELM、error_signatures", "需要读取计算目录输出"],
    ], [1700, 2600, 2800, 2260])
    add_h2(doc, "1.2 当前本机环境")
    add_bullets(doc, [
        "已登记 7 台服务器；截图时 cl9 与 minus 处于连接状态，其余服务器离线。",
        "默认服务器是 cl9；cl9 使用 PBS，minus 使用 Slurm，其余服务器由 auto 探测。",
        "已配置 Codex、GLM 与 DeepSeek 等 4 个模型 Provider；最近缓存探测均为 full。",
        "本地 Web UI 源码版本为 1.3.5；Codex 插件实测优先加载已安装的 1.3.0 Client 后端。",
        "当前配置的服务器仍为 interactive 模式；密钥免交互重连代码已具备，但尚未逐台启用。",
    ])

    add_h1(doc, "2. 快速开始")
    add_h2(doc, "2.1 使用前检查")
    add_bullets(doc, [
        "Windows 已安装 Python 3.11 或更高版本，推荐 Python 3.12。",
        "系统能够使用 ssh 与 scp，并已正确登记 Vlab 的主机指纹。",
        "已配置 Vlab 身份文件路径；身份文件内容不应粘贴到聊天或界面文本框。",
        "Vlab 上已部署 vaspilot-gateway；各 HPC 服务器已登记名称、user@host、根目录和调度器。",
        "至少配置一个可用模型 Provider；没有模型也可以使用文件、作业、监测和 CLI 功能。",
    ])
    add_h2(doc, "2.2 启动本地 Web 控制台")
    add_code(doc, "scripts\\vaspilot.cmd ui\n"
                  "# 不自动打开浏览器\n"
                  "scripts\\vaspilot.cmd ui --no-open")
    add_body(doc, "启动后，终端会打印带随机会话令牌的本地地址，例如 http://127.0.0.1:8930/t/<token>。每次重新启动都会生成新令牌，旧标签页出现“会话已过期”属于正常现象。若 8930 端口被占用，服务会自动尝试相邻端口。")
    add_h2(doc, "2.3 第一次连接服务器")
    add_numbers(doc, [
        "在左侧服务器列表找到目标服务器，确认服务器名称和调度器标签。",
        "单击“连接”。交互认证服务器会打开一个独立、可见的系统终端。",
        "在终端中输入服务器密码和当前 TOTP 验证码；密码与验证码不会进入模型上下文。",
        "返回 Web UI，点击服务器区域的“刷新”，确认状态点由红色变为绿色。",
        "进入“监测”页面检查该服务器是否能返回作业列表和调度器类型。",
    ])
    add_note(doc, "不要混淆", "Web UI 中的连接按钮不接受密码。若对话模型询问密码、验证码或私钥正文，应停止该流程并改用可见终端。", "danger")

    add_h1(doc, "3. Codex 插件与集群面板")
    add_figure(doc, CODEX_DASHBOARD,
               "图 1  Codex 内嵌远程集群作业面板",
               "Codex embedded dashboard with server cards, active jobs and progress inputs")
    add_h2(doc, "3.1 在 Codex 中打开面板")
    add_code(doc, "显示 VASPilot 集群作业面板，并每 30 秒刷新一次。\n"
                  "只显示 cl9 和 minus 的集群面板。\n"
                  "查询 cl9 上 /public/home/.../runs/case-1 的 VASP 科学进度。")
    add_h2(doc, "3.2 界面区域")
    add_table(doc, ["区域", "用途", "如何判断"], [
        ["顶部 KPI", "在线服务器、当天作业、活动服务器", "用于快速判断集群整体是否可用"],
        ["LIVE 与刷新", "显示轮询状态并手动立即刷新", "默认 30 秒；允许 10–300 秒"],
        ["服务器卡片", "连接状态、登录目标、调度器与作业摘要", "绿色在线；红色显示离线原因"],
        ["运行中/排队", "当前调度器活动作业数量", "来自 Slurm/PBS，不能代替科学判断"],
        ["科学进度", "输入计算目录后查询 OSZICAR/OUTCAR", "返回离子步、电子步、能量和收敛标志"],
    ], [1700, 4300, 3360])
    add_h2(doc, "3.3 面板能做什么")
    add_bullets(doc, [
        "一次查看全部或指定服务器的连接状态、调度器、活动作业和作业状态。",
        "对用户指定且位于 remote_root 内的计算目录查询 VASP 科学进度。",
        "对 key 模式服务器执行无人值守重连，并区分退避、密钥拒绝、网络不可达和主机指纹异常。",
        "在不支持 MCP App 渲染的宿主中，退化为结构化文本结果。",
    ])
    add_h2(doc, "3.4 面板不能替代什么")
    add_bullets(doc, [
        "它不是调度器的推送订阅，而是定时轮询。关闭 Codex 页面后不会继续在页面内刷新。",
        "它不会自动知道每个 Job ID 对应哪个 VASP 目录；科学进度通常需要明确目录。",
        "COMPLETED 只表示调度器结束；需进一步读取 VASP 科学进度。",
        "面板默认设计为查询视图；提交、审批和密钥安装应在明确的工作流或可见终端中完成。",
    ])

    add_h1(doc, "4. Web 控制台总览与对话")
    add_figure(doc, UI / "ui-01-chat.png",
               "图 2  Web 控制台“对话”页：项目、服务器、模型与对话区",
               "VASPilot web console chat page with project list, server list and model selector")
    add_h2(doc, "4.1 固定导航")
    add_bullets(doc, [
        "顶部标签：对话、文件、作业、工作流、监测、配置。",
        "左侧项目区：新建、固定、选择和删除本地计算项目。",
        "左侧服务器区：显示连接点、调度器、默认服务器、连接/断开按钮。",
        "左下模型区：选择模型 Provider，查看 full 或 analysis_only 模式，并重新进行能力探测。",
    ])
    add_h2(doc, "4.2 对话页", page_break=True)
    add_numbers(doc, [
        "先选择与任务相关的本地项目；这样模型能够读取项目文件并把下载结果放入受限项目目录。",
        "选择模型。full 模式可调用写工具；analysis_only 仅允许读取与分析。",
        "用自然语言描述目标、服务器和路径；涉及科学结论时要求模型同时报告调度器状态与科学状态。",
        "观察工具调用摘要。展开工具结果时重点检查服务器名、路径、Job ID 和返回状态。",
        "对话会按会话保存；新建任务时可新开会话，避免不同计算上下文混在一起。",
    ])
    add_note(doc, "提示词写法", "把服务器名、远端绝对路径、希望得到的证据和允许的动作写清楚。例如：“只读检查 cl9 上 /public/home/.../case-1，报告作业状态、最后能量和是否收敛，不修改文件。”", "ok")

    add_h1(doc, "5. 服务器连接与认证")
    add_h2(doc, "5.1 首次配置 Vlab 中转密钥")
    add_body(doc, "首次使用时，先在 Vlab 网站创建一对 SSH 密钥并把私钥文件保存到本机。这个私钥用于“本机 → Vlab 网关”的连接；完成一次配置后，VASPilot 才能通过 Vlab 再访问已登记的远端服务器。")
    add_h3(doc, "步骤 1：进入 Vlab 的 SSH 密钥管理")
    add_body(doc, "登录 https://vlab.ustc.edu.cn/vm/，在顶部点击“虚拟机管理”，再点击“SSH 密钥管理”。")
    add_figure(doc, VLAB / "vlab-01-home.png",
               "步骤图 A  Vlab 门户：从“虚拟机管理”进入 SSH 密钥管理",
               "Vlab portal navigation showing Virtual Machine Management")
    add_figure(doc, VLAB / "vlab-02-key-management.png",
               "步骤图 B  点击“SSH 密钥管理”",
               "Vlab SSH key management button", width=3.1)
    add_h3(doc, "步骤 2：生成并下载私钥")
    add_body(doc, "在密钥管理页面点击“生成新的 SSH 密钥对”。生成完成后，点击“下载私钥”，将下载的文件保存到本机固定位置；后续配置填写的是这个文件的完整路径，而不是把私钥文本粘贴到网页或聊天框。")
    add_figure(doc, VLAB / "vlab-03-create-download.png",
               "步骤图 C  生成新的 SSH 密钥对，并下载私钥文件",
               "Vlab page showing Generate new SSH key pair and Download private key", width=2.75)
    add_h3(doc, "步骤 3：在 VASPilot 中填写连接设置")
    add_numbers(doc, [
        "在 VASPilot Web UI 顶部进入“配置”，选择“连接设置”。",
        "在“Vlab 主机”填入 vlab.ustc.edu.cn；如果 Vlab 或管理员另行提供 SSH 主机名、端口，以提供的信息为准。",
        "填写你的 Vlab 用户名和 SSH 端口；未另行说明时端口通常为 22。",
        "在“身份文件路径”选择刚才下载的私钥文件，例如 C:\\Users\\<用户名>\\.ssh\\vaspilot_vlab_key。",
        "保存后执行一次连接检查；随后再连接任一已登记的 HPC 服务器。",
    ])
    add_note(doc, "配置要点", "这里填写的是本机私钥文件的路径。不要在“身份文件路径”中粘贴私钥内容，也不要把 Vlab 网站地址误填成远端 HPC 服务器地址。", "ok")
    add_h2(doc, "5.2 交互认证模式")
    add_body(doc, "interactive 是当前 7 台服务器的实际配置。没有可复用 SSH 会话时，查询会返回 auth_required 或 disconnected；用户必须通过 Web UI 的“连接”、Codex 的 open_remote_login，或 CLI 的 server connect 打开可见终端。")
    add_code(doc, "scripts\\vaspilot.cmd server list\n"
                  "scripts\\vaspilot.cmd server status cl9\n"
                  "scripts\\vaspilot.cmd server connect cl9\n"
                  "scripts\\vaspilot.cmd server disconnect cl9")
    add_h2(doc, "5.3 每服务器密钥模式")
    add_body(doc, "密钥模式会在 Vlab 上为每台 HPC 生成独立 Ed25519 密钥，将公钥安装到目标服务器并验证 BatchMode 登录。验证成功后设置 auth_mode=key、auto_connect=true；会话丢失时按 30、60、120、300 秒逐级退避重连。")
    add_code(doc, "scripts\\vaspilot.cmd server key-setup cl9\n"
                  "scripts\\vaspilot.cmd server key-status cl9\n"
                  "scripts\\vaspilot.cmd server key-disable cl9\n"
                  "scripts\\vaspilot.cmd server key-revoke cl9 --confirm-server cl9")
    add_note(doc, "启用条件", "只有服务器端 SSH 策略允许公钥登录绕过密码/TOTP，并且 BatchMode 验证成功，免交互重连才成立。主机指纹变化时系统会停止自动重试。", "warn")
    add_h2(doc, "5.4 离线状态解释")
    add_table(doc, ["界面状态", "常见原因", "处理方式"], [
        ["需要登录", "interactive 会话不存在或已过期", "打开可见终端重新输入密码与 TOTP"],
        ["自动重连中", "key 模式正在尝试建立 ControlMaster", "等待本次尝试完成"],
        ["等待退避", "连续失败，尚未到下一次重试时间", "等待倒计时或人工执行 server connect"],
        ["公钥被拒绝", "authorized_keys 缺失、权限或策略变化", "重新执行 key-setup 或恢复交互模式"],
        ["网络不可达", "DNS、链路、防火墙或服务器离线", "先检查网络与 Vlab，再检查 HPC"],
        ["主机指纹异常", "known_hosts 与目标指纹不一致", "人工核实管理员公告；不要自动接受新指纹"],
    ], [1900, 3600, 3860], compact=True)

    add_h1(doc, "6. 项目与远端文件")
    add_figure(doc, UI / "ui-02-files.png",
               "图 3  “文件”页：选择服务器、输入根目录内路径并浏览",
               "VASPilot files page showing server selector, remote path and directory table")
    add_h2(doc, "6.1 本地项目")
    add_body(doc, "本地项目位于 VASPilot 配置目录下，主要保存 INCAR、KPOINTS、POSCAR、可选 run.job.sh，以及远端 POTCAR 库位置。项目可以与对话会话绑定，使模型在多轮任务中保持同一个计算上下文。")
    add_bullets(doc, [
        "INCAR、KPOINTS、POSCAR 可在项目编辑器中创建和修改。",
        "POTCAR 正文不通过模型展示；界面仅记录元数据或远端库位置。",
        "项目校验会检查必要文件、INCAR 基础参数、KPOINTS 和 POSCAR 的可解析性。",
        "上传和下载都绑定项目根目录，目标已存在时默认拒绝覆盖。",
    ])
    add_h2(doc, "6.2 文件页操作")
    add_numbers(doc, [
        "选择服务器。界面会优先填入配置的 remote_root；若未配置，则查询登录主目录。",
        "输入远端绝对路径并单击“浏览”。路径必须位于该服务器允许根目录内。",
        "单击目录继续下钻；使用“上一级”返回。",
        "单击文本文件查看内容。POTCAR、WAVECAR、CHGCAR 等大文件或敏感科学文件会拒绝文本读取。",
        "单击“测量全部”时会递归计算当前目录各项真实大小；大型目录可能耗时。",
    ])
    add_h2(doc, "6.3 文件操作语义")
    add_table(doc, ["操作", "语义", "关键限制"], [
        ["list/read/tail/find/stat/du", "只读查看、搜索与统计", "路径限制在 remote_root；文本读取有大小和文件类型限制"],
        ["mkdir", "创建目录", "不得越出根目录"],
        ["copy/move", "同一服务器内复制或移动", "源和目标都必须在同一服务器根目录"],
        ["upload/download", "本机与远端之间传输单个文件", "SHA-256 校验；默认不覆盖"],
        ["trash/restore", "移入可恢复回收站或恢复", "避免直接永久删除"],
        ["purge", "永久删除回收站条目", "必须重复输入完全相同的 trash ID"],
    ], [1650, 3200, 4510])
    add_code(doc, "scripts\\vaspilot.cmd remote list --server cl9 /public/home/...\n"
                  "scripts\\vaspilot.cmd remote tail --server cl9 /public/home/.../OUTCAR --lines 120\n"
                  "scripts\\vaspilot.cmd remote du --server cl9 /public/home/.../case-1\n"
                  "scripts\\vaspilot.cmd remote trash-list --server cl9")

    add_h1(doc, "7. 作业状态与 VASP 科学进度")
    add_figure(doc, UI / "ui-03-jobs.png",
               "图 4  “作业”页：活动作业、历史台账和独立的 VASP 科学进度查询",
               "VASPilot jobs page with active jobs, job history and VASP progress query")
    add_h2(doc, "7.1 活动作业与历史作业")
    add_body(doc, "选择服务器后单击“查询作业”。活动区显示当前排队或运行的作业；历史区合并调度器历史与本机台账。部分完成时间属于“首次发现作业消失”的推断值，界面或数据中会以 assumed_end 标识。")
    add_h2(doc, "7.2 常见调度器状态")
    add_table(doc, ["状态", "含义", "下一步"], [
        ["PENDING", "正在排队或等待资源", "检查队列、分区、资源和优先级"],
        ["RUNNING", "已分配资源并运行", "结合 elapsed 与 VASP 进度判断是否正常"],
        ["COMPLETING", "调度器正在收尾", "等待终态，不要提前判断科学结果"],
        ["COMPLETED", "调度器认为作业正常结束", "立即查询 VASP 科学进度"],
        ["FAILED/TIMEOUT", "程序失败或超过时限", "查看 OUTCAR、作业输出和诊断签名"],
        ["OUT_OF_MEMORY", "内存不足", "检查资源配置和体系规模"],
        ["CANCELLED/PREEMPTED", "人工取消或被抢占", "确认是否需要新尝试"],
        ["UNKNOWN", "调度器无记录或解析不足", "交叉检查 recent、输出文件和 Job ID"],
    ], [1750, 3900, 3710])
    add_h2(doc, "7.3 VASP 科学进度")
    add_numbers(doc, [
        "选择与作业相同的服务器。",
        "输入 VASP 计算目录，而不是 OUTCAR 文件本身。",
        "单击“查询”，读取 INCAR、OSZICAR 和 OUTCAR 的受限内容。",
        "检查 ionic_steps、last_ionic、electronic_reached_nelm、error_signatures。",
        "只有 scientific_converged=true 才能描述为科学收敛；若调度器完成但科学未收敛，应标记为 needs_review。",
    ])
    add_code(doc, "scripts\\vaspilot.cmd job list --server cl9\n"
                  "scripts\\vaspilot.cmd job recent --server cl9\n"
                  "scripts\\vaspilot.cmd job progress --server cl9 /public/home/.../case-1\n"
                  "scripts\\vaspilot.cmd job diagnose --server cl9 /public/home/.../case-1")

    add_h1(doc, "8. 工作流：计划、审批、运行与恢复")
    add_figure(doc, UI / "ui-04-workflow.png",
               "图 5  “工作流”页第一步：选择本地项目、服务器、远端目录和资源参数",
               "VASPilot workflow page with local project, remote directory and scheduler parameters")
    add_h2(doc, "8.1 计划包含什么")
    add_bullets(doc, [
        "目标服务器、远端目录、调度器、作业名、分区、核数、时限与 VASP 可执行文件。",
        "每个输入文件的本地路径、远端路径、大小和 SHA-256。",
        "上传、校验、提交、监测、科学进度、下载和解析组成的步骤 DAG。",
        "风险摘要：创建哪些目录、是否触碰调度器、覆盖策略和资源参数。",
    ])
    add_h2(doc, "8.2 标准四步")
    add_numbers(doc, [
        "准备计划：选择本地项目或输入目录，填写服务器、远端目录和资源参数，单击“生成预览”。",
        "检查预览：逐项检查文件哈希、远端路径、作业脚本、步骤 DAG 和风险摘要。",
        "本地审批：输入界面要求的确认短语。审批令牌绑定服务器、计划哈希和文件哈希，并有有效期。",
        "运行与监测：执行同一不可变计划；每一步结果写入运行状态。失败后的 resume 会创建新尝试，并保留旧证据。",
    ])
    add_code(doc, "scripts\\vaspilot.cmd workflow prepare `\n"
                  "  --from-dir C:\\calc\\case-01 --server cl9 `\n"
                  "  --remote-dir /public/home/.../runs/case-01 `\n"
                  "  --scheduler pbs --job-name case01 --ntasks 8 --walltime 24:00:00\n\n"
                  "scripts\\vaspilot.cmd workflow preview <plan_id>\n"
                  "scripts\\vaspilot.cmd workflow approve <plan_id>\n"
                  "scripts\\vaspilot.cmd workflow run <plan_id> --approval-ref <ref>\n"
                  "scripts\\vaspilot.cmd workflow status <plan_id>")
    add_note(doc, "不可变计划", "审批后若服务器、文件内容、作业脚本或步骤发生变化，plan_hash 会改变，原审批失效。上传前还会再次核对本地文件哈希。", "ok")
    add_h2(doc, "8.3 当前工作流能力边界")
    add_body(doc, "可执行引擎目前实现的是通用单阶段计划：mkdir、upload、validate、submit、monitor、progress、download、parse。仓库中的 relax、DOS、NEB、电荷差和收敛扫描 JSON 是设计配方，尚未由当前引擎自动加载执行。复杂多阶段任务仍需智能体通过通用工具逐阶段组织，或等待配方执行器接入。")

    add_h1(doc, "9. 集群监测与资源视图")
    add_figure(doc, UI / "ui-05-monitor.png",
               "图 6  “监测”页总览：连接数、活动作业与各服务器调度器快照",
               "VASPilot monitoring page with cluster KPI cards and scheduler snapshot")
    add_h2(doc, "9.1 总览")
    add_bullets(doc, [
        "顶部 KPI 显示已连接服务器数、活动作业总数、存在排队或运行作业的服务器数。",
        "集群快照按服务器显示连接、调度器、活动作业和状态集合。",
        "勾选自动刷新后，仅在监测页面可见时每 60 秒刷新一次，避免无意义请求。",
        "该页面适合回答“哪些服务器在线”“哪里有空闲”“哪些集群正在排队”。",
    ])
    add_h2(doc, "9.2 空闲算力与单机详情")
    add_body(doc, "服务器资源卡会读取 CPU、负载、内存、磁盘、GPU 利用率、显存、温度、功耗和调度器队列。单击卡片进入单机详情，可查看 GPU 进程的用户归属、资源热力图和 GPU 使用分布。无 GPU 的服务器仍可显示 CPU、内存和磁盘信息。")
    add_h2(doc, "9.3 离线采集器")
    add_body(doc, "可选离线采集器安装在服务器 remote_root 下的 .vp-monitor 目录。Web UI 在线读取时会写入心跳；心跳超过 90 秒后，远端守护循环每分钟采样一次，从而补齐关闭浏览器期间的资源历史。历史 TSV 受大小上限约束并会定期裁剪。")
    add_code(doc, "scripts\\vaspilot.cmd monitor snapshot\n"
                  "scripts\\vaspilot.cmd monitor watch --servers cl9,minus --interval 30")

    add_h1(doc, "10. 配置模型、连接、联网搜索和技能")
    add_figure(doc, UI / "ui-06-settings.png",
               "图 7  “配置”页的模型服务区；API Key 只显示保存状态，不回显明文",
               "VASPilot settings page showing model provider cards without plaintext API keys")
    add_h2(doc, "10.1 模型 Provider")
    add_table(doc, ["协议", "适用后端", "特点"], [
        ["openai-chat-compatible", "GLM、DeepSeek、Ollama、LM Studio 等", "Chat Completions 风格流式工具调用"],
        ["openai-responses", "OpenAI Responses API", "Responses 事件流与函数工具"],
        ["codex-sdk", "Codex SDK 或本机 codex CLI", "SDK 优先，CLI 作为后备"],
    ], [2300, 3300, 3760])
    add_numbers(doc, [
        "填写显示名、协议、API 地址、模型名和可选环境变量名。",
        "将 API Key 保存到 Windows DPAPI 凭据存储；界面保存后不会回显明文。",
        "单击“检测”验证可达性、流式输出和工具调用能力。",
        "探测通过后模式为 full；失败时降级为 analysis_only。",
        "确认目标模型后单击“设为默认”，再回到对话页选择。",
    ])
    add_h2(doc, "10.2 连接设置")
    add_body(doc, "连接设置保存 Vlab 主机、用户、端口、身份文件路径和网关程序路径。首次配置的完整点击流程见 5.1；通常填写 Vlab 主机 vlab.ustc.edu.cn、你的 Vlab 用户名、平台要求的端口，以及刚下载的本机私钥文件绝对路径。只保存路径与非敏感元数据，不应把 PEM 内容、服务器密码或 TOTP 种子写入配置。")
    add_h2(doc, "10.3 智能体执行策略")
    add_table(doc, ["设置", "作用", "注意"], [
        ["confirm", "job_submit 先生成待确认卡片", "适合日常团队使用；远端脚本哈希变化会阻止提交"],
        ["auto", "job_submit 直接进入调度器", "当前本机配置为 auto；不会出现提交确认卡片"],
        ["最大工具轮次", "限制一次智能体任务的工具循环", "达到软上限后模型会被提醒收尾；仍有硬上限"],
    ], [1700, 3600, 4060])
    add_h2(doc, "10.4 联网搜索与技能")
    add_bullets(doc, [
        "联网搜索支持智谱、博查与免密钥 Bing SERP；Bing 依赖页面结构，改版时可能返回空结果。",
        "web_fetch 只允许公开 HTTP(S) 地址，拒绝回环、内网地址和异常端口，并限制返回大小。",
        "技能用于保存可复用的领域说明与操作纪律；技能文本会进入模型系统上下文，但不能替代后端权限检查。",
    ])

    add_h1(doc, "11. CLI 命令参考")
    add_body(doc, "开发环境建议始终使用仓库自带包装器 scripts\\vaspilot.cmd，避免系统中其他旧版本 vaspilot 抢占入口。每个普通命令输出一个稳定 JSON 文档。")
    add_table(doc, ["命令组", "主要子命令", "用途"], [
        ["server", "list/add/edit/remove/connect/disconnect/status/doctor/key-*", "服务器目录、连接、诊断与密钥模式"],
        ["remote", "pwd/list/read/tail/find/stat/du/upload/download/mkdir/copy/move/trash/restore/purge", "受根目录约束的远端文件操作"],
        ["job", "list/recent/submit/cancel/progress/diagnose", "调度器状态和 VASP 科学检查"],
        ["workflow", "prepare/validate/preview/approve/run/resume/status", "不可变计划与审批执行"],
        ["monitor", "snapshot/watch", "多服务器快照和变化检测"],
        ["agent", "provider/chat/run", "配置模型、普通对话与目标驱动工具循环"],
        ["ui", "ui", "启动本地 Web 控制台"],
    ], [1450, 4850, 3060])
    add_h2(doc, "11.1 退出码")
    add_table(doc, ["退出码", "含义", "典型场景"], [
        ["0", "成功", "命令完成并输出 ok"],
        ["1", "一般错误", "远端程序或未分类异常"],
        ["2", "用法错误", "命令组、参数或格式错误"],
        ["3", "需要认证", "Vlab 或 HPC 可复用会话不存在"],
        ["4", "需要/无效审批", "审批缺失、过期、绑定不符或重复使用"],
        ["5", "验证失败", "路径、Job ID、输入文件或参数未通过校验"],
    ], [1350, 2800, 5210])
    add_h2(doc, "11.2 常用只读命令")
    add_code(doc, "scripts\\vaspilot.cmd server list\n"
                  "scripts\\vaspilot.cmd monitor snapshot\n"
                  "scripts\\vaspilot.cmd job list --server cl9\n"
                  "scripts\\vaspilot.cmd job recent --server cl9\n"
                  "scripts\\vaspilot.cmd remote list --server cl9 /public/home/...\n"
                  "scripts\\vaspilot.cmd remote tail --server cl9 /public/home/.../OUTCAR --lines 120\n"
                  "scripts\\vaspilot.cmd job progress --server cl9 /public/home/.../case-1")
    add_h2(doc, "11.3 取消和永久删除")
    add_body(doc, "取消作业和 purge 都采用双重匹配参数，避免模型或用户误选对象。操作前先重新查询目标状态，复制准确 Job ID 或 Trash ID，再重复输入确认值。")

    add_h1(doc, "12. 四个典型使用流程")
    add_h2(doc, "12.1 每日集群巡检")
    add_numbers(doc, [
        "在 Codex 中打开集群面板，或进入 Web UI“监测”。",
        "确认在线服务器数量与预期一致；记录离线分类。",
        "查看活动作业数量和 PENDING/RUNNING 分布。",
        "进入资源详情，比较 CPU、GPU、显存与队列拥堵情况。",
        "对关键作业输入计算目录，确认 VASP 科学进度。",
    ])
    add_h2(doc, "12.2 检查已有 VASP 计算")
    add_numbers(doc, [
        "在“文件”页确认目录存在，检查 INCAR、OSZICAR、OUTCAR 与 CONTCAR。",
        "在“作业”页查询对应 Job ID 的调度器状态。",
        "用科学进度查询读取离子步、最后能量、NELM 与错误签名。",
        "若调度器仍运行但能量或电子步异常，先保存证据再提出恢复方案。",
        "若调度器完成但 scientific_converged=false，将结果标记为需要复核。",
    ])
    add_h2(doc, "12.3 多服务器选择")
    add_numbers(doc, [
        "先用监测页比较服务器在线状态、调度器、空闲资源和队列。",
        "确认目标服务器上存在需要的软件模块、VASP 可执行文件和 POTCAR 库。",
        "明确 PBS 队列或 Slurm 分区；不要依赖模型猜测默认队列。",
        "在生成计划前固定目标服务器；更换服务器需要重新生成和审批计划。",
    ])
    add_h2(doc, "12.4 从本地项目提交一次标准计算", page_break=True)
    add_numbers(doc, [
        "新建项目并写入 INCAR、KPOINTS、POSCAR；登记远端 POTCAR 库位置。",
        "执行项目校验，修复明确错误；不要把 POTCAR 正文交给模型。",
        "在工作流页填写服务器、远端目录、核数、分区和时限，生成计划。",
        "检查计划哈希、文件哈希、作业脚本、DAG 和风险摘要。",
        "完成审批后运行；跟踪 submit、monitor、progress、download 和 parse 各步骤。",
        "最终同时记录 Job ID、调度器终态、科学收敛、结果文件哈希和本地下载位置。",
    ])

    add_h1(doc, "13. 故障排查")
    add_table(doc, ["现象", "判断路径", "处理"], [
        ["auth_required", "检查 Vlab 与目标服务器会话", "打开连接终端重新认证；不要把密码发给模型"],
        ["UI 会话已过期", "本地服务已重启，旧 token 失效", "重新运行 vaspilot ui，使用新地址"],
        ["端口被占用", "启动日志显示实际端口", "使用日志中的新 URL；无需手工杀进程"],
        ["gateway unavailable", "server doctor 检查身份文件、网关路径和版本", "重新部署网关，再做只读自检"],
        ["host_key_failed", "目标主机指纹变化", "人工向管理员核实；确认前停止连接"],
        ["path outside root", "路径不在 remote_root 或包含 ..", "从 remote_pwd 返回的根目录重新构造绝对路径"],
        ["Job ID 查不到", "active 与 recent 都无记录", "核对服务器、调度器与 Job ID；检查本机台账和输出文件"],
        ["作业完成但未收敛", "scientific_converged=false", "查看 NELM、错误签名和最后能量，创建新恢复尝试"],
        ["Provider analysis_only", "能力探测未通过", "检查 API Key、地址、模型与工具调用能力后重新探测"],
        ["Bing 返回空结果", "SERP 页面结构或网络变化", "改用智谱/博查，或直接 fetch 明确公开 URL"],
    ], [2500, 3100, 3760])
    add_h2(doc, "13.1 诊断原则")
    add_bullets(doc, [
        "先确定问题属于连接、调度器、VASP 科学状态、文件边界还是模型 Provider。",
        "先做只读检查，再决定是否需要写操作或新尝试。",
        "保留失败尝试、日志、Job ID、计划哈希和输入文件哈希。",
        "恢复任务创建新尝试，不覆盖旧目录或旧证据。",
    ])

    add_h1(doc, "14. 权限、审计与高权限工具")
    add_h2(doc, "14.1 命名工具的边界")
    add_bullets(doc, [
        "服务器名、远端路径、Job ID、Trash ID 和本地项目路径均经过格式验证。",
        "named remote 工具会在本地和 Vlab 网关两侧检查 remote_root。",
        "上传与下载记录 SHA-256；下载目标已存在时拒绝覆盖。",
        "POTCAR 只返回 TITEL、ENMAX、大小和哈希等元数据。",
        "审计日志按 UTC 日期追加写入，并对疑似密码、Token 和 API Key 字段脱敏。",
    ])
    add_h2(doc, "14.2 实际存在的高权限工具")
    add_note(doc, "真实实现",
             "当前 full 模式注册表包含 shell_run 与 remote_run。它们允许执行任意本地或远端命令，属于“全程审计、事后追踪”而不是“确定性拦截”。它们可以绕过命名工具的路径边界和工作流审批。",
             "danger")
    add_bullets(doc, [
        "只读或生产场景应优先使用 remote_*、job_*、vasp_* 和 workflow_* 命名工具。",
        "analysis_only 模式会阻止 shell_run、remote_run 和其他写工具。",
        "当前 MCP 会把共享注册表工具原样导出给 Codex；插件说明中的“无 Shell”与实际代码不一致。",
        "agent_submit_mode=auto 时，job_submit 没有人工确认步骤；需要确认卡片时应切回 confirm。",
    ])
    add_h2(doc, "14.3 审批工作流仍然提供的保证")
    add_body(doc, "在按 workflow prepare/approve/run 使用时，审批令牌由本地 HMAC 密钥签名，绑定服务器、plan_hash、files_hash、动作和参数哈希，并采用一次性消费语义。该保证只覆盖通过工作流引擎执行的计划，不能约束另行执行的任意 Shell。")

    add_h1(doc, "15. 已知限制与版本管理")
    add_bullets(doc, [
        "监控是轮询，不是调度器推送；刷新间隔内可能短暂显示旧状态。",
        "高级 VASP Recipe 尚未接入当前执行引擎。",
        "真实集群验收仍取决于各服务器的队列、模块、VASP 安装、授权目录与 SSH 策略。",
        "无密钥 Bing 搜索依赖搜索结果 HTML 结构，稳定性低于正式 API。",
        "Codex 插件当前按发现顺序优先加载 %LOCALAPPDATA% 中的 1.3.0 Client 后端，而不是 1.3.5 源码仓库。",
        "源码 __version__、pyproject、MCP Server、插件和已安装包元数据目前不是同一个版本号。",
        "测试覆盖 288 个离线测试，但仓库暂时没有 CI，也没有在本手册编制过程中提交真实作业。",
    ])
    add_h2(doc, "15.1 判断正在使用哪个后端")
    add_code(doc, "scripts\\vaspilot.cmd --version\n"
                  "py -3.12 C:\\Users\\weikx\\plugins\\vaspilot-remote-control\\scripts\\vaspilot_remote_mcp.py --self-test\n"
                  "# Codex 中也可调用 vaspilot_self_check")
    add_body(doc, "插件后端发现顺序为：VASPILOT_ROOT → 显式 Client 配置 → %LOCALAPPDATA%\\VASPilot\\client.json → 插件本地配置 → 打包 backend → 开发路径。排查“源码已有功能但 Codex 看不到”时，首先确认插件实际加载的 backend_root。")

    add_h1(doc, "附录 A：状态词典")
    add_table(doc, ["字段/术语", "说明"], [
        ["connected", "存在可用的服务器 SSH 复用会话"],
        ["auth_mode", "interactive 或 key"],
        ["auto_connect", "是否允许 key 模式自动恢复会话"],
        ["scheduler_state", "Slurm/PBS 生命周期状态"],
        ["scientific_converged", "VASP 科学收敛总判断"],
        ["ionic_converged", "离子弛豫是否达到要求；NSW=0 静态任务按电子结果处理"],
        ["electronic_reached_nelm", "电子迭代是否触及 NELM 上限"],
        ["error_signatures", "从 OUTCAR 识别的已知严重错误"],
        ["plan_hash", "不可变计划内容的 SHA-256"],
        ["approval_ref", "绑定计划和参数的一次性审批引用"],
        ["assumed_end", "本机台账根据作业从调度器历史中消失而推断结束"],
    ], [2600, 6760])

    add_h1(doc, "附录 B：可直接使用的提示词")
    add_h2(doc, "只读巡检")
    add_code(doc, "只读检查所有已登记服务器：报告连接状态、调度器、活动作业数量和状态。\n"
                  "不要修改文件，不要提交或取消作业。对离线服务器说明具体离线分类。")
    add_h2(doc, "科学进度")
    add_code(doc, "检查 cl9 上 /public/home/.../runs/case-1：\n"
                  "1. 查询调度器状态；2. 查询 VASP 科学进度；3. 报告最后离子步与能量；\n"
                  "4. 明确区分 scheduler_state 与 scientific_converged；5. 不做任何修改。")
    add_h2(doc, "多服务器选型")
    add_code(doc, "比较 cl9、minus 和 cl12 的连接、队列、CPU/GPU/内存与活动作业。\n"
                  "只给出适合提交本次任务的候选服务器和依据；不要实际提交。")
    add_h2(doc, "准备工作流")
    add_code(doc, "为本地项目 case-01 生成到 cl9 的 VASP 工作流预览。\n"
                  "目标目录 /public/home/.../runs/case-01，8 核，24 小时。\n"
                  "展示所有文件哈希、作业脚本、步骤 DAG 和风险；停在审批前。")

    add_h1(doc, "附录 C：新成员交接检查表")
    add_bullets(doc, [
        "能够启动 Web UI，并理解带 token 的本地 URL。",
        "能够在可见终端完成 interactive 登录，不把密码或 TOTP 发给模型。",
        "能够区分服务器连接、调度器状态和 VASP 科学状态。",
        "能够浏览 remote_root 内目录并读取 OUTCAR/OSZICAR 的受限内容。",
        "能够查询 active、recent 和科学进度，并解释 COMPLETED 的边界。",
        "能够生成并检查工作流计划，理解 plan_hash 和文件 SHA-256。",
        "知道 confirm 与 auto 提交策略的差异。",
        "知道 shell_run/remote_run 是高权限、审计型工具，不等同于命名工具边界。",
        "知道 Codex 插件可能加载已安装 Client 后端，而非当前源码目录。",
        "发生失败时保留原目录、日志、Job ID 和计划证据，再创建恢复尝试。",
    ])

    # Final paragraph helps establish Bottom bookmark in viewers if later added.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("— 文稿结束 —")
    set_run_font(r, size=9, color=MUTED, italic=True)

    patch_numbering(doc)
    doc.core_properties.title = "VASPilot 详细使用手册"
    doc.core_properties.subject = "Codex 插件、本地 Web 控制台与 CLI 使用指南"
    doc.core_properties.author = "VASPilot 项目组"
    doc.core_properties.keywords = "VASPilot, VASP, HPC, Codex, Slurm, PBS, Vlab"
    doc.core_properties.comments = "基于 VASPilot 1.3.5 源码界面编制"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
